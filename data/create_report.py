#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
APA 7 Formatında Kapsamlı Analiz Raporu Oluşturucu
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os

# Veri dosyalarını oku
data_path = "/home/user/gamlj_docs/data/"

# CSV dosyalarını oku
descriptives = pd.read_csv(os.path.join(data_path, "descriptives_python.csv"), index_col=0)
correlations = pd.read_csv(os.path.join(data_path, "significant_correlations_python.csv"))
bayes_means = pd.read_csv(os.path.join(data_path, "bayes_bootstrap_means.csv"))
bayes_corr = pd.read_csv(os.path.join(data_path, "bayes_bootstrap_correlations.csv"))
bayes_reg = pd.read_csv(os.path.join(data_path, "bayes_bootstrap_regression.csv"))
gender_ttest = pd.read_csv(os.path.join(data_path, "gender_ttest_results.csv"))
gender_desc = pd.read_csv(os.path.join(data_path, "gender_descriptives.csv"))

# Word belgesi oluştur
doc = Document()

# Stil ayarları
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Başlık stili
title_style = doc.styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
title_style.font.name = 'Times New Roman'
title_style.font.size = Pt(14)
title_style.font.bold = True

# Heading stili
heading_style = doc.styles.add_style('HeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
heading_style.font.name = 'Times New Roman'
heading_style.font.size = Pt(12)
heading_style.font.bold = True

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    p.space_after = Pt(12)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_before = Pt(12)
    p.space_after = Pt(6)

def add_paragraph(text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    p.space_after = Pt(0)

def add_table_note(text):
    p = doc.add_paragraph()
    run = p.add_run("Not. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.space_before = Pt(6)

# =============================================
# RAPOR İÇERİĞİ
# =============================================

# BAŞLIK
add_title("Psikolojik Sıkıntı Değişkenlerinin Network Psikometri Yaklaşımıyla İncelenmesi:")
add_title("Bayesian Bootstrap Analizi")
doc.add_paragraph()

# Yazar bilgileri
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Yazar Adı]")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Kurum Adı]")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# ÖZET
add_heading("Özet")
add_paragraph("""Bu araştırma, psikolojik sıkıntı göstergeleri (depresyon, anksiyete, stres) ile öz-şefkat, utanç, suçluluk, sosyal medya bağımlılığı, sosyal destek ve siber mağduriyet arasındaki ilişkileri incelemeyi amaçlamaktadır. Kesitsel tarama deseniyle yürütülen çalışmada 1028 Türk yetişkin katılımcıdan veri toplanmıştır. Veriler network psikometri yaklaşımı ve Bayesian bootstrap analizi kullanılarak değerlendirilmiştir. Bulgular, depresyon, anksiyete ve stres arasında güçlü pozitif ilişkiler olduğunu göstermiştir. Öz-şefkat ile psikolojik sıkıntı değişkenleri arasında orta düzeyde pozitif ilişkiler saptanmıştır. Aile ve arkadaş desteğinin psikolojik sıkıntı ile negatif ilişkili olduğu bulunmuştur. Cinsiyet karşılaştırmalarında kadınların erkeklere göre daha yüksek depresyon, anksiyete ve stres düzeyleri gösterdiği tespit edilmiştir. Sonuçlar, psikolojik müdahalelerde sosyal destek ve öz-şefkat odaklı yaklaşımların önemini vurgulamaktadır.""", indent=True)

p = doc.add_paragraph()
run = p.add_run("Anahtar Kelimeler: ")
run.font.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run = p.add_run("psikolojik sıkıntı, depresyon, anksiyete, stres, network analizi, Bayesian bootstrap")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# GİRİŞ
add_heading("Giriş")
add_paragraph("""Psikolojik sıkıntı, bireylerin günlük yaşam işlevselliğini olumsuz etkileyen depresyon, anksiyete ve stres gibi belirtileri kapsayan geniş bir kavramdır. Modern yaşamın getirdiği stres faktörleri, özellikle sosyal medya kullanımı ve siber zorbalık deneyimleri, bireylerin psikolojik iyilik halini tehdit etmektedir. Bu bağlamda, psikolojik sıkıntının koruyucu ve risk faktörlerinin anlaşılması büyük önem taşımaktadır.""", indent=True)

add_paragraph("""Network psikometri yaklaşımı, psikolojik yapılar arasındaki karmaşık ilişkileri görselleştirmek ve analiz etmek için güçlü bir çerçeve sunmaktadır. Bu yaklaşım, geleneksel faktör analitiği modellerden farklı olarak, değişkenlerin birbirleriyle doğrudan etkileşim içinde olduğunu varsaymaktadır. Bu sayede, psikolojik semptomlar arasındaki dinamik ilişkiler daha iyi anlaşılabilmektedir.""", indent=True)

add_paragraph("""Bu araştırmanın amacı, psikolojik sıkıntı göstergeleri (depresyon, anksiyete, stres) ile utanç, suçluluk, öz-şefkat, sosyal medya bağımlılığı, sosyal destek ve siber mağduriyet arasındaki ilişkileri network psikometri yaklaşımıyla incelemektir. Ayrıca, bu ilişkilerin cinsiyete göre farklılaşıp farklılaşmadığı değerlendirilmiştir.""", indent=True)

doc.add_page_break()

# YÖNTEM
add_heading("Yöntem")

add_heading("Araştırma Deseni", level=2)
add_paragraph("""Bu araştırma, nicel araştırma yöntemlerinden kesitsel tarama deseni kullanılarak yürütülmüştür. Çalışmada network psikometri yaklaşımı benimsenmiş ve değişkenler arasındaki ilişkiler network analizi aracılığıyla incelenmiştir.""", indent=True)

add_heading("Katılımcılar", level=2)
add_paragraph("""Çalışmaya toplam 1028 Türk yetişkin katılmıştır. Katılımcıların yaş ortalaması 31.30 (SS = 9.34) olup, yaş aralığı 18-67 arasında değişmektedir. Örneklemin cinsiyet dağılımı ve diğer demografik özellikleri Tablo 1'de sunulmaktadır.""", indent=True)

add_heading("Veri Toplama Araçları", level=2)
add_paragraph("""Depresyon, Anksiyete ve Stres Ölçeği-21 (DASS-21). Psikolojik sıkıntı düzeylerini ölçmek için DASS-21 kullanılmıştır. Ölçek 21 maddeden oluşmakta ve her alt ölçek (depresyon, anksiyete, stres) 7 madde içermektedir.""", indent=True)

add_paragraph("""Öz-Şefkat Ölçeği. Katılımcıların öz-şefkat düzeylerini ölçmek için 10 maddelik Öz-Şefkat Ölçeği kullanılmıştır.""", indent=True)

add_paragraph("""Bergen Sosyal Medya Bağımlılığı Ölçeği. Sosyal medya bağımlılığı düzeylerini değerlendirmek için 6 maddelik Bergen Sosyal Medya Bağımlılığı Ölçeği kullanılmıştır.""", indent=True)

add_paragraph("""Sosyal Destek Ölçeği. Aile (4 madde) ve arkadaş (4 madde) sosyal desteğini ölçmek için 8 maddelik Sosyal Destek Ölçeği kullanılmıştır.""", indent=True)

add_paragraph("""Utanç ve Suçluluk Ölçeği. Utanç (5 madde) ve suçluluk (5 madde) düzeylerini değerlendirmek için 11 maddelik Utanç ve Suçluluk Ölçeği kullanılmıştır.""", indent=True)

add_paragraph("""Siber Mağduriyet Ölçeği. Siber mağduriyet deneyimlerini ölçmek için 10 maddelik Siber Mağduriyet Ölçeği kullanılmıştır.""", indent=True)

add_heading("Veri Analizi", level=2)
add_paragraph("""Veri analizi sürecinde R ve Python programlama dilleri kullanılmıştır. Betimleyici istatistikler (ortalama, standart sapma, çarpıklık, basıklık) hesaplanmıştır. Değişkenler arası ilişkiler Pearson korelasyon analizi ile incelenmiştir. Bayesian bootstrap analizi, parametre tahminlerinin belirsizliğini değerlendirmek ve güvenilir aralıklar oluşturmak için uygulanmıştır. Cinsiyet grupları arasındaki farklılıklar bağımsız örneklemler t-testi ve Cohen's d etki büyüklüğü ile değerlendirilmiştir.""", indent=True)

doc.add_page_break()

# BULGULAR
add_heading("Bulgular")

add_heading("Betimleyici İstatistikler", level=2)
add_paragraph("""Araştırma değişkenlerine ilişkin betimleyici istatistikler Tablo 1'de sunulmaktadır. Depresyon ortalaması 0.97 (SS = 0.67), anksiyete ortalaması 0.73 (SS = 0.60) ve stres ortalaması 1.07 (SS = 0.66) olarak bulunmuştur. Tüm psikolojik sıkıntı değişkenleri pozitif çarpıklık göstermiştir.""", indent=True)

# Tablo 1: Betimleyici İstatistikler
p = doc.add_paragraph()
run = p.add_run("Tablo 1")
run.font.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Araştırma Değişkenlerine İlişkin Betimleyici İstatistikler")
run.font.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Tablo oluştur
table1 = doc.add_table(rows=1, cols=7)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Başlık satırı
header_cells = table1.rows[0].cells
headers = ['Değişken', 'N', 'Ort.', 'SS', 'Min', 'Maks', 'Çarpıklık']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

# Veri satırları
var_names_tr = {
    'Depression': 'Depresyon',
    'Anxiety': 'Anksiyete',
    'Stress': 'Stres',
    'Shame': 'Utanç',
    'Guilt': 'Suçluluk',
    'SelfCompassion': 'Öz-Şefkat',
    'SocialMediaAddiction': 'Sosyal Medya Bağ.',
    'FamilySupport': 'Aile Desteği',
    'FriendSupport': 'Arkadaş Desteği',
    'CyberVictimization': 'Siber Mağduriyet',
    'SocialMediaTime': 'Sosyal Medya Süresi',
    'Age': 'Yaş'
}

for var in descriptives.index:
    row = table1.add_row().cells
    row[0].text = var_names_tr.get(var, var)
    row[1].text = str(int(descriptives.loc[var, 'count']))
    row[2].text = f"{descriptives.loc[var, 'mean']:.2f}"
    row[3].text = f"{descriptives.loc[var, 'std']:.2f}"
    row[4].text = f"{descriptives.loc[var, 'min']:.2f}"
    row[5].text = f"{descriptives.loc[var, 'max']:.2f}"
    row[6].text = f"{descriptives.loc[var, 'skewness']:.2f}"
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_table_note("Ort. = Ortalama; SS = Standart Sapma; Sosyal Medya Bağ. = Sosyal Medya Bağımlılığı")

doc.add_paragraph()

# Korelasyon Analizi Sonuçları
add_heading("Korelasyon Analizi Sonuçları", level=2)
add_paragraph("""Değişkenler arası Pearson korelasyon analizi sonuçları incelendiğinde, depresyon ile stres arasında güçlü pozitif ilişki (r = .82, p < .001) bulunmuştur. Benzer şekilde, anksiyete ile stres (r = .79, p < .001) ve depresyon ile anksiyete (r = .74, p < .001) arasında da güçlü pozitif ilişkiler saptanmıştır. Öz-şefkat ile psikolojik sıkıntı değişkenleri arasında orta düzeyde pozitif ilişkiler bulunmuştur: stres (r = .68, p < .001), depresyon (r = .66, p < .001) ve anksiyete (r = .65, p < .001). Aile desteği ile depresyon (r = -.36, p < .001), stres (r = -.26, p < .001) ve anksiyete (r = -.25, p < .001) arasında negatif ilişkiler saptanmıştır.""", indent=True)

# Tablo 2: Önemli Korelasyonlar
p = doc.add_paragraph()
run = p.add_run("Tablo 2")
run.font.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Değişkenler Arası Anlamlı Korelasyonlar (En Güçlü 15 İlişki)")
run.font.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells = table2.rows[0].cells
headers = ['Değişken 1', 'Değişken 2', 'r', 'p']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

for idx, corr_row in correlations.head(15).iterrows():
    row = table2.add_row().cells
    row[0].text = var_names_tr.get(corr_row['From'], corr_row['From'])
    row[1].text = var_names_tr.get(corr_row['To'], corr_row['To'])
    row[2].text = f"{corr_row['Correlation']:.3f}"
    row[3].text = "< .001"
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_table_note("r = Pearson korelasyon katsayısı")

doc.add_paragraph()
doc.add_page_break()

# Bayesian Bootstrap Analizi Sonuçları
add_heading("Bayesian Bootstrap Analizi Sonuçları", level=2)
add_paragraph("""Bayesian bootstrap analizi ile elde edilen parametre tahminleri ve güven aralıkları Tablo 3'te sunulmaktadır. Bu analiz, geleneksel istatistiksel yöntemlere kıyasla daha sağlam ve yorumlanması kolay sonuçlar sunmaktadır.""", indent=True)

# Tablo 3: Bayesian Bootstrap Ortalamaları
p = doc.add_paragraph()
run = p.add_run("Tablo 3")
run.font.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Bayesian Bootstrap Analizi ile Elde Edilen Ortalama ve Güven Aralıkları")
run.font.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

table3 = doc.add_table(rows=1, cols=6)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells = table3.rows[0].cells
headers = ['Değişken', 'Ort.', 'SS', '%95 GA Alt', '%95 GA Üst', 'HDI Alt']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

for idx, row_data in bayes_means.iterrows():
    row = table3.add_row().cells
    row[0].text = var_names_tr.get(row_data['Variable'], row_data['Variable'])
    row[1].text = f"{row_data['Mean']:.3f}"
    row[2].text = f"{row_data['SD']:.3f}"
    row[3].text = f"{row_data['CI_Lower']:.3f}"
    row[4].text = f"{row_data['CI_Upper']:.3f}"
    row[5].text = f"{row_data['HDI_Lower']:.3f}"
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_table_note("Ort. = Ortalama; SS = Standart Sapma; GA = Güven Aralığı; HDI = Highest Density Interval")

doc.add_paragraph()

# Bayesian Bootstrap Regresyon Sonuçları
add_heading("Depresyon Yordayıcıları: Bayesian Bootstrap Regresyon Analizi", level=2)
add_paragraph("""Depresyonu yordayan değişkenlerin Bayesian bootstrap regresyon analizi sonuçları Tablo 4'te sunulmaktadır. Sonuçlar, stresin (β = 0.56, %95 GA [0.49, 0.62]) ve anksiyetenin (β = 0.22, %95 GA [0.15, 0.29]) depresyonun en güçlü yordayıcıları olduğunu göstermektedir. Öz-şefkat (β = 0.08, %95 GA [0.05, 0.11]) ve suçluluk (β = 0.07, %95 GA [0.02, 0.11]) da anlamlı yordayıcılar olarak bulunmuştur.""", indent=True)

# Tablo 4: Regresyon Sonuçları
p = doc.add_paragraph()
run = p.add_run("Tablo 4")
run.font.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Depresyonun Bayesian Bootstrap Regresyon Analizi Sonuçları")
run.font.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

table4 = doc.add_table(rows=1, cols=6)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells = table4.rows[0].cells
headers = ['Yordayıcı', 'β', 'SS', '%95 GA Alt', '%95 GA Üst', 'P(+)']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

predictor_names_tr = {
    'Intercept': 'Sabit',
    'Anxiety': 'Anksiyete',
    'Stress': 'Stres',
    'Shame': 'Utanç',
    'Guilt': 'Suçluluk',
    'SelfCompassion': 'Öz-Şefkat',
    'SocialMediaAddiction': 'Sosyal Medya Bağ.',
    'CyberVictimization': 'Siber Mağduriyet'
}

for idx, row_data in bayes_reg.iterrows():
    row = table4.add_row().cells
    row[0].text = predictor_names_tr.get(row_data['Predictor'], row_data['Predictor'])
    row[1].text = f"{row_data['Coefficient']:.3f}"
    row[2].text = f"{row_data['SD']:.3f}"
    row[3].text = f"{row_data['CI_Lower']:.3f}"
    row[4].text = f"{row_data['CI_Upper']:.3f}"
    row[5].text = f"{row_data['Prob_Positive']:.3f}"
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_table_note("β = Standardize edilmemiş regresyon katsayısı; SS = Standart Sapma; GA = Güven Aralığı; P(+) = Katsayının pozitif olma olasılığı; Sosyal Medya Bağ. = Sosyal Medya Bağımlılığı")

doc.add_paragraph()
doc.add_page_break()

# Cinsiyet Karşılaştırması
add_heading("Cinsiyet Grupları Karşılaştırması", level=2)
add_paragraph("""Cinsiyet grupları arasındaki farklılıklar bağımsız örneklemler t-testi ile incelenmiştir (Tablo 5). Sonuçlar, kadınların erkeklere göre anlamlı düzeyde daha yüksek depresyon (t = -2.48, p = .013, d = 0.18), anksiyete (t = -2.60, p = .009, d = 0.19) ve stres (t = -3.03, p = .003, d = 0.22) puanlarına sahip olduğunu göstermiştir. Ayrıca, kadınların öz-şefkat puanları da erkeklerden anlamlı düzeyde daha yüksek bulunmuştur (t = -2.62, p = .009, d = 0.19). Öte yandan, erkeklerin siber mağduriyet puanları kadınlardan anlamlı düzeyde daha yüksektir (t = 3.37, p < .001, d = -0.31).""", indent=True)

# Tablo 5: Cinsiyet Karşılaştırması
p = doc.add_paragraph()
run = p.add_run("Tablo 5")
run.font.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Cinsiyet Gruplarına Göre Değişkenlerin Karşılaştırılması")
run.font.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

table5 = doc.add_table(rows=1, cols=8)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells = table5.rows[0].cells
headers = ['Değişken', 'Erkek Ort.', 'Erkek SS', 'Kadın Ort.', 'Kadın SS', 't', 'p', "Cohen's d"]
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)

for idx, row_data in gender_ttest.iterrows():
    desc_row = gender_desc[gender_desc['Variable'] == row_data['Variable']]
    if len(desc_row) > 0:
        row = table5.add_row().cells
        row[0].text = var_names_tr.get(row_data['Variable'], row_data['Variable'])
        row[1].text = f"{desc_row['Male_Mean'].values[0]:.2f}"
        row[2].text = f"{desc_row['Male_SD'].values[0]:.2f}"
        row[3].text = f"{desc_row['Female_Mean'].values[0]:.2f}"
        row[4].text = f"{desc_row['Female_SD'].values[0]:.2f}"
        row[5].text = f"{row_data['t']:.2f}"
        p_val = row_data['p']
        sig = row_data['Sig'] if pd.notna(row_data['Sig']) else ''
        row[6].text = f"{p_val:.3f}{sig}"
        cohens_d = row_data["Cohen's d"]
        row[7].text = f"{cohens_d:.2f}"
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)

add_table_note("Ort. = Ortalama; SS = Standart Sapma; *p < .05, **p < .01, ***p < .001")

doc.add_paragraph()
doc.add_page_break()

# TARTIŞMA
add_heading("Tartışma")
add_paragraph("""Bu araştırmada, psikolojik sıkıntı göstergeleri ile çeşitli psikososyal değişkenler arasındaki ilişkiler network psikometri yaklaşımı ve Bayesian bootstrap analizi kullanılarak incelenmiştir. Bulgular, depresyon, anksiyete ve stres arasında güçlü pozitif ilişkiler olduğunu göstermiştir. Bu sonuç, psikolojik sıkıntının birbirleriyle yakından ilişkili belirtiler kümesi olarak kavramlaştırılmasını desteklemektedir.""", indent=True)

add_paragraph("""Öz-şefkat ile psikolojik sıkıntı değişkenleri arasında bulunan orta düzeyde pozitif ilişkiler dikkat çekicidir. Bu bulgu, bireylerin kendilerine karşı şefkatli tutumlarının psikolojik sıkıntı yaşama eğilimleriyle ilişkili olabileceğini göstermektedir. Ancak, bu ilişkinin yönü ve nedenselliği konusunda ileri araştırmalara ihtiyaç duyulmaktadır.""", indent=True)

add_paragraph("""Aile ve arkadaş sosyal desteğinin psikolojik sıkıntı ile negatif ilişkili olması, sosyal desteğin koruyucu bir faktör olarak işlev gördüğünü desteklemektedir. Bu bulgu, psikolojik müdahalelerde sosyal destek sistemlerinin güçlendirilmesinin önemini vurgulamaktadır.""", indent=True)

add_paragraph("""Cinsiyet karşılaştırmalarında kadınların erkeklere göre daha yüksek depresyon, anksiyete ve stres düzeyleri göstermesi, literatürle tutarlıdır. Öte yandan, erkeklerin daha yüksek siber mağduriyet deneyimlerine sahip olması dikkat çekici bir bulgudur ve bu konuda daha fazla araştırmaya ihtiyaç vardır.""", indent=True)

add_paragraph("""Bayesian bootstrap regresyon analizi, depresyonun en güçlü yordayıcılarının stres ve anksiyete olduğunu göstermiştir. Bu bulgu, psikolojik sıkıntı belirtilerinin birbirini tetikleyen bir yapıya sahip olduğunu desteklemektedir. Klinik uygulamalar açısından, bir belirtinin tedavisi diğer belirtilerin azalmasına da katkıda bulunabilir.""", indent=True)

# SONUÇ
add_heading("Sonuç")
add_paragraph("""Bu araştırma, psikolojik sıkıntı değişkenlerinin karmaşık ilişki yapısını network psikometri yaklaşımıyla ortaya koymaktadır. Bulgular, psikolojik müdahalelerde sosyal destek ve öz-şefkat odaklı yaklaşımların önemini vurgulamaktadır. Cinsiyet farklılıkları, müdahale programlarının cinsiyete duyarlı olarak tasarlanması gerektiğini göstermektedir. Gelecek araştırmaların boylamsal desenlerle nedensellik ilişkilerini incelemesi önerilmektedir.""", indent=True)

doc.add_page_break()

# KAYNAKÇA
add_heading("Kaynakça")
add_paragraph("[Kaynakça APA 7 formatında eklenecektir]")

# Belgeyi kaydet
output_path = os.path.join(data_path, "Analiz_Raporu_Kapsamli_APA7.docx")
doc.save(output_path)
print(f"Rapor başarıyla oluşturuldu: {output_path}")
