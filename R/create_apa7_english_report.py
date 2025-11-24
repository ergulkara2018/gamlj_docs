#!/usr/bin/env python3
"""
APA 7th Edition Format - Analysis Report (English)
Network Analysis of Social Media Use and Psychological Distress
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os

# Change to data directory
os.chdir('/home/user/gamlj_docs')

# =============================================================================
# LOAD DATA
# =============================================================================

descriptives = pd.read_csv('data/descriptives_python.csv', index_col=0)
bayes_means = pd.read_csv('data/bayes_bootstrap_means.csv')
bayes_regression = pd.read_csv('data/bayes_bootstrap_regression.csv')
gender_ttest = pd.read_csv('data/gender_ttest_results.csv')
gender_desc = pd.read_csv('data/gender_descriptives.csv')
centrality = pd.read_csv('data/centrality_python.csv')
correlations = pd.read_csv('data/significant_correlations_python.csv')

# =============================================================================
# CREATE DOCUMENT
# =============================================================================

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins (1 inch = 2.54 cm)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# =============================================================================
# TITLE PAGE
# =============================================================================

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Network Analysis of Social Media Use, Cyber Victimization, and Psychological Distress: A Bayesian Bootstrap Approach')
title_run.bold = True
title_run.font.size = Pt(14)

doc.add_paragraph()  # Empty line

# Authors (placeholder)
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run('Author Name').italic = True

doc.add_paragraph()

# Affiliation
affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil.add_run('University Affiliation').italic = True

doc.add_paragraph()
doc.add_paragraph()

# =============================================================================
# ABSTRACT
# =============================================================================

abstract_heading = doc.add_paragraph()
abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract_run = abstract_heading.add_run('Abstract')
abstract_run.bold = True

abstract_text = doc.add_paragraph()
abstract_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstract_text.add_run(
    'The present study investigated the relationships among social media use, cyber victimization, '
    'shame, guilt, self-compassion, social support, and psychological distress (depression, anxiety, and stress) '
    'using network analysis and Bayesian bootstrap methods. A total of 1,028 Turkish participants '
    '(predominantly female) completed self-report measures assessing these constructs. Network analysis '
    'revealed that depression, anxiety, and stress were highly interconnected and served as central nodes '
    'in the psychological network. Bayesian bootstrap regression indicated that stress (b = 0.56, 95% CI [0.49, 0.62]) '
    'and anxiety (b = 0.22, 95% CI [0.15, 0.29]) were the strongest predictors of depression. '
    'Gender comparison analyses revealed significant differences, with females reporting higher levels of '
    'depression, anxiety, stress, and self-compassion, while males reported higher cyber victimization. '
    'These findings highlight the complex interplay between digital experiences and mental health outcomes.'
)

# Keywords
keywords = doc.add_paragraph()
keywords_label = keywords.add_run('Keywords: ')
keywords_label.italic = True
keywords.add_run('network analysis, psychological distress, social media, cyber victimization, Bayesian bootstrap')

doc.add_page_break()

# =============================================================================
# INTRODUCTION
# =============================================================================

intro_heading = doc.add_paragraph()
intro_run = intro_heading.add_run('Introduction')
intro_run.bold = True

intro_p1 = doc.add_paragraph()
intro_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro_p1.paragraph_format.first_line_indent = Inches(0.5)
intro_p1.add_run(
    'Social media use has become an integral part of daily life, particularly among young adults. '
    'While these platforms offer opportunities for social connection and information sharing, '
    'increasing evidence suggests potential links between intensive social media use and adverse '
    'mental health outcomes. The present study employed network psychometrics and Bayesian bootstrap '
    'methods to examine the complex relationships among social media use, cyber victimization, '
    'self-conscious emotions, and psychological distress.'
)

intro_p2 = doc.add_paragraph()
intro_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro_p2.paragraph_format.first_line_indent = Inches(0.5)
intro_p2.add_run(
    'Network analysis provides a novel framework for understanding psychological phenomena by '
    'conceptualizing mental health constructs as interconnected networks of symptoms and behaviors '
    'rather than latent variables. This approach allows for identification of central nodes that '
    'may serve as intervention targets and reveals the dynamic relationships among psychological variables.'
)

# =============================================================================
# METHOD
# =============================================================================

method_heading = doc.add_paragraph()
method_run = method_heading.add_run('Method')
method_run.bold = True

# Participants
participants_heading = doc.add_paragraph()
participants_run = participants_heading.add_run('Participants')
participants_run.bold = True
participants_run.italic = True

participants_p = doc.add_paragraph()
participants_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
participants_p.paragraph_format.first_line_indent = Inches(0.5)
participants_p.add_run(
    f'The sample consisted of 1,028 Turkish participants (M_age = 31.30, SD = 9.34, range = 18-67). '
    f'Participants were recruited through convenience sampling methods. The study was approved by '
    f'the institutional ethics committee, and all participants provided informed consent prior to participation.'
)

# Measures
measures_heading = doc.add_paragraph()
measures_run = measures_heading.add_run('Measures')
measures_run.bold = True
measures_run.italic = True

# DASS-21
dass_p = doc.add_paragraph()
dass_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
dass_p.paragraph_format.first_line_indent = Inches(0.5)
dass_label = dass_p.add_run('Depression Anxiety Stress Scales-21 (DASS-21). ')
dass_label.bold = True
dass_p.add_run(
    'The DASS-21 (Lovibond & Lovibond, 1995) is a 21-item self-report measure assessing depression '
    '(7 items), anxiety (7 items), and stress (7 items) over the past week. Items are rated on a '
    '4-point scale ranging from 0 (did not apply to me at all) to 3 (applied to me very much or most of the time).'
)

# Shame and Guilt
sg_p = doc.add_paragraph()
sg_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sg_p.paragraph_format.first_line_indent = Inches(0.5)
sg_label = sg_p.add_run('Shame and Guilt Scale. ')
sg_label.bold = True
sg_p.add_run(
    'Shame (5 items) and guilt (5 items) were assessed using the Shame and Guilt Scale. '
    'This measure distinguishes between these self-conscious emotions and their differential effects on psychological well-being.'
)

# Self-Compassion
sc_p = doc.add_paragraph()
sc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sc_p.paragraph_format.first_line_indent = Inches(0.5)
sc_label = sc_p.add_run('Self-Compassion Scale. ')
sc_label.bold = True
sc_p.add_run(
    'The 10-item Self-Compassion Scale was used to measure participants\' capacity to be kind, '
    'understanding, and non-judgmental toward themselves during difficult times.'
)

# Bergen
bergen_p = doc.add_paragraph()
bergen_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
bergen_p.paragraph_format.first_line_indent = Inches(0.5)
bergen_label = bergen_p.add_run('Bergen Social Media Addiction Scale. ')
bergen_label.bold = True
bergen_p.add_run(
    'The 6-item Bergen Social Media Addiction Scale was used to assess problematic social media use patterns, '
    'including mood modification, tolerance, withdrawal, conflict, and relapse.'
)

# Social Support
ss_p = doc.add_paragraph()
ss_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ss_p.paragraph_format.first_line_indent = Inches(0.5)
ss_label = ss_p.add_run('Social Support Scale. ')
ss_label.bold = True
ss_p.add_run(
    'The 8-item Social Support Scale measured family support (4 items) and friend support (4 items).'
)

# Cyber Victimization
cv_p = doc.add_paragraph()
cv_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
cv_p.paragraph_format.first_line_indent = Inches(0.5)
cv_label = cv_p.add_run('Cyber Victimization Scale. ')
cv_label.bold = True
cv_p.add_run(
    'The 10-item Cyber Victimization Scale assessed various forms of online victimization '
    'including harassment, exclusion, and threatening behaviors.'
)

# Data Analysis
analysis_heading = doc.add_paragraph()
analysis_run = analysis_heading.add_run('Data Analysis')
analysis_run.bold = True
analysis_run.italic = True

analysis_p = doc.add_paragraph()
analysis_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
analysis_p.paragraph_format.first_line_indent = Inches(0.5)
analysis_p.add_run(
    'Data were analyzed using Python 3.x with the following packages: pandas for data manipulation, '
    'scipy for statistical tests, networkx for network analysis, and pgmpy for Bayesian network estimation. '
    'Bayesian bootstrap analysis (Rubin, 1981) with 4,000 iterations was employed to obtain robust '
    'estimates with 95% credible intervals. Network centrality measures (degree, betweenness, closeness, '
    'and strength) were calculated to identify influential nodes. Gender comparisons were conducted using '
    'independent samples t-tests with Welch\'s correction when assumptions of homogeneity of variance were violated.'
)

doc.add_page_break()

# =============================================================================
# RESULTS
# =============================================================================

results_heading = doc.add_paragraph()
results_run = results_heading.add_run('Results')
results_run.bold = True

# Descriptive Statistics
desc_heading = doc.add_paragraph()
desc_run = desc_heading.add_run('Descriptive Statistics')
desc_run.bold = True
desc_run.italic = True

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
desc_p.paragraph_format.first_line_indent = Inches(0.5)
desc_p.add_run(
    'Table 1 presents the descriptive statistics for all study variables. Depression, anxiety, and stress '
    'scores were in the normal to mild range. Cyber victimization was relatively low (M = 1.17, SD = 0.35), '
    'indicating infrequent experiences of online victimization in the sample.'
)

# Table 1 - Descriptive Statistics
table1_caption = doc.add_paragraph()
table1_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
table1_caption_run = table1_caption.add_run('Table 1')
table1_caption_run.italic = True
table1_caption.add_run('\n')
table1_caption.add_run('Descriptive Statistics for Study Variables (N = 1,028)').italic = True

# Create table
table1 = doc.add_table(rows=1, cols=6)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
header_cells = table1.rows[0].cells
headers = ['Variable', 'M', 'SD', 'Skewness', 'Kurtosis', '95% CI']
for i, header in enumerate(headers):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].runs[0].font.bold = True
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
variable_labels = {
    'Depression': 'Depression',
    'Anxiety': 'Anxiety',
    'Stress': 'Stress',
    'Shame': 'Shame',
    'Guilt': 'Guilt',
    'SelfCompassion': 'Self-Compassion',
    'SocialMediaAddiction': 'Social Media Addiction',
    'FamilySupport': 'Family Support',
    'FriendSupport': 'Friend Support',
    'CyberVictimization': 'Cyber Victimization'
}

for var in ['Depression', 'Anxiety', 'Stress', 'Shame', 'Guilt', 'SelfCompassion',
            'SocialMediaAddiction', 'FamilySupport', 'FriendSupport', 'CyberVictimization']:
    row = table1.add_row()
    bayes_row = bayes_means[bayes_means['Variable'] == var].iloc[0]
    desc_row = descriptives.loc[var]

    row.cells[0].text = variable_labels.get(var, var)
    row.cells[1].text = f"{bayes_row['Mean']:.2f}"
    row.cells[2].text = f"{bayes_row['SD']:.2f}"
    row.cells[3].text = f"{desc_row['skewness']:.2f}"
    row.cells[4].text = f"{desc_row['kurtosis']:.2f}"
    row.cells[5].text = f"[{bayes_row['CI_Lower']:.2f}, {bayes_row['CI_Upper']:.2f}]"

    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Table note
table1_note = doc.add_paragraph()
table1_note.add_run('Note. ').italic = True
table1_note.add_run('M = mean; SD = standard deviation; CI = credible interval from Bayesian bootstrap (4,000 iterations).')

doc.add_paragraph()

# Network Analysis Results
network_heading = doc.add_paragraph()
network_run = network_heading.add_run('Network Analysis')
network_run.bold = True
network_run.italic = True

network_p = doc.add_paragraph()
network_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
network_p.paragraph_format.first_line_indent = Inches(0.5)
network_p.add_run(
    'Network analysis revealed a highly interconnected structure among psychological distress variables. '
    'The strongest correlations were observed between depression and stress (r = .82, p < .001), '
    'anxiety and stress (r = .79, p < .001), and depression and anxiety (r = .74, p < .001). '
    'Self-compassion showed strong positive correlations with all three distress measures, suggesting '
    'potential maladaptive self-compassion in this sample.'
)

# Centrality Analysis
centrality_heading = doc.add_paragraph()
centrality_run = centrality_heading.add_run('Centrality Analysis')
centrality_run.bold = True
centrality_run.italic = True

centrality_p = doc.add_paragraph()
centrality_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
centrality_p.paragraph_format.first_line_indent = Inches(0.5)
centrality_p.add_run(
    'Centrality analyses (Table 2) identified depression, stress, and anxiety as the most central nodes '
    'in the network based on strength centrality. Social media addiction demonstrated the highest '
    'betweenness centrality, suggesting its role as a bridge connecting different clusters of variables.'
)

# Table 2 - Centrality
table2_caption = doc.add_paragraph()
table2_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
table2_caption_run = table2_caption.add_run('Table 2')
table2_caption_run.italic = True
table2_caption.add_run('\n')
table2_caption.add_run('Network Centrality Measures').italic = True

table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells2 = table2.rows[0].cells
headers2 = ['Variable', 'Degree', 'Betweenness', 'Closeness', 'Strength']
for i, header in enumerate(headers2):
    header_cells2[i].text = header
    header_cells2[i].paragraphs[0].runs[0].font.bold = True
    header_cells2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for _, cent_row in centrality.iterrows():
    row = table2.add_row()
    row.cells[0].text = variable_labels.get(cent_row['Variable'], cent_row['Variable'])
    row.cells[1].text = f"{cent_row['Degree']:.2f}"
    row.cells[2].text = f"{cent_row['Betweenness']:.2f}"
    row.cells[3].text = f"{cent_row['Closeness']:.2f}"
    row.cells[4].text = f"{cent_row['Strength']:.2f}"

    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Bayesian Bootstrap Regression
regression_heading = doc.add_paragraph()
regression_run = regression_heading.add_run('Bayesian Bootstrap Regression Analysis')
regression_run.bold = True
regression_run.italic = True

regression_p = doc.add_paragraph()
regression_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
regression_p.paragraph_format.first_line_indent = Inches(0.5)
regression_p.add_run(
    'Table 3 presents the results of Bayesian bootstrap regression predicting depression. '
    'Stress emerged as the strongest predictor (b = 0.56, 95% CI [0.49, 0.62], P(b > 0) = 1.00), '
    'followed by anxiety (b = 0.22, 95% CI [0.15, 0.29], P(b > 0) = 1.00). Self-compassion also '
    'showed a significant positive association with depression (b = 0.08, 95% CI [0.05, 0.11], P(b > 0) = 1.00). '
    'Guilt was a significant predictor (b = 0.07, 95% CI [0.02, 0.11], P(b > 0) = .996), while shame '
    'and cyber victimization did not reach significance.'
)

# Table 3 - Regression
table3_caption = doc.add_paragraph()
table3_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
table3_caption_run = table3_caption.add_run('Table 3')
table3_caption_run.italic = True
table3_caption.add_run('\n')
table3_caption.add_run('Bayesian Bootstrap Regression Predicting Depression').italic = True

table3 = doc.add_table(rows=1, cols=5)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells3 = table3.rows[0].cells
headers3 = ['Predictor', 'b', 'SD', '95% CI', 'P(b > 0)']
for i, header in enumerate(headers3):
    header_cells3[i].text = header
    header_cells3[i].paragraphs[0].runs[0].font.bold = True
    header_cells3[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

predictor_labels = {
    'Intercept': 'Intercept',
    'Anxiety': 'Anxiety',
    'Stress': 'Stress',
    'Shame': 'Shame',
    'Guilt': 'Guilt',
    'SelfCompassion': 'Self-Compassion',
    'SocialMediaAddiction': 'Social Media Addiction',
    'CyberVictimization': 'Cyber Victimization'
}

for _, reg_row in bayes_regression.iterrows():
    row = table3.add_row()
    row.cells[0].text = predictor_labels.get(reg_row['Predictor'], reg_row['Predictor'])
    row.cells[1].text = f"{reg_row['Coefficient']:.3f}"
    row.cells[2].text = f"{reg_row['SD']:.3f}"
    row.cells[3].text = f"[{reg_row['CI_Lower']:.3f}, {reg_row['CI_Upper']:.3f}]"
    row.cells[4].text = f"{reg_row['Prob_Positive']:.3f}"

    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

table3_note = doc.add_paragraph()
table3_note.add_run('Note. ').italic = True
table3_note.add_run('Based on 4,000 Bayesian bootstrap iterations. P(b > 0) = posterior probability that the coefficient is positive.')

doc.add_paragraph()

# Gender Comparison
gender_heading = doc.add_paragraph()
gender_run = gender_heading.add_run('Gender Comparison Analysis')
gender_run.bold = True
gender_run.italic = True

gender_p = doc.add_paragraph()
gender_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
gender_p.paragraph_format.first_line_indent = Inches(0.5)
gender_p.add_run(
    'Table 4 presents the results of gender comparison analyses. Females reported significantly higher '
    'levels of depression (t = -2.48, p = .013, d = 0.18), anxiety (t = -2.60, p = .009, d = 0.19), '
    'stress (t = -3.03, p = .003, d = 0.22), and self-compassion (t = -2.62, p = .009, d = 0.19). '
    'In contrast, males reported significantly higher levels of cyber victimization '
    '(t = 3.37, p < .001, d = -0.31). Effect sizes ranged from small to medium.'
)

# Table 4 - Gender Comparison
table4_caption = doc.add_paragraph()
table4_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
table4_caption_run = table4_caption.add_run('Table 4')
table4_caption_run.italic = True
table4_caption.add_run('\n')
table4_caption.add_run('Gender Comparison: Descriptive Statistics and Independent Samples t-Tests').italic = True

table4 = doc.add_table(rows=1, cols=7)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells4 = table4.rows[0].cells
headers4 = ['Variable', 'Male M (SD)', 'Female M (SD)', 't', 'p', 'd', 'Effect Size']
for i, header in enumerate(headers4):
    header_cells4[i].text = header
    header_cells4[i].paragraphs[0].runs[0].font.bold = True
    header_cells4[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

effect_labels = {
    'Kucuk': 'Small',
    'Orta': 'Medium',
    'Buyuk': 'Large',
    'Cok Buyuk': 'Very Large'
}

for _, ttest_row in gender_ttest.iterrows():
    var = ttest_row['Variable']
    desc_row_g = gender_desc[gender_desc['Variable'] == var].iloc[0]

    row = table4.add_row()
    row.cells[0].text = variable_labels.get(var, var)
    row.cells[1].text = f"{desc_row_g['Male_Mean']:.2f} ({desc_row_g['Male_SD']:.2f})"
    row.cells[2].text = f"{desc_row_g['Female_Mean']:.2f} ({desc_row_g['Female_SD']:.2f})"
    row.cells[3].text = f"{ttest_row['t']:.2f}"

    p_val = ttest_row['p']
    if p_val < 0.001:
        row.cells[4].text = "< .001***"
    elif p_val < 0.01:
        row.cells[4].text = f".{str(p_val)[2:4]}**"
    elif p_val < 0.05:
        row.cells[4].text = f".{str(p_val)[2:4]}*"
    else:
        row.cells[4].text = f".{str(p_val)[2:4]}"

    cohens_d_val = ttest_row["Cohen's d"]
    row.cells[5].text = f"{cohens_d_val:.2f}"
    row.cells[6].text = effect_labels.get(ttest_row['Effect'], ttest_row['Effect'])

    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

table4_note = doc.add_paragraph()
table4_note.add_run('Note. ').italic = True
table4_note.add_run('d = Cohen\'s d. *p < .05. **p < .01. ***p < .001.')

doc.add_page_break()

# =============================================================================
# DISCUSSION
# =============================================================================

discussion_heading = doc.add_paragraph()
discussion_run = discussion_heading.add_run('Discussion')
discussion_run.bold = True

discussion_p1 = doc.add_paragraph()
discussion_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
discussion_p1.paragraph_format.first_line_indent = Inches(0.5)
discussion_p1.add_run(
    'The present study employed network analysis and Bayesian bootstrap methods to examine the relationships '
    'among social media use, cyber victimization, self-conscious emotions, and psychological distress. '
    'Several key findings emerged from the analyses.'
)

discussion_p2 = doc.add_paragraph()
discussion_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
discussion_p2.paragraph_format.first_line_indent = Inches(0.5)
discussion_p2.add_run(
    'First, network analysis revealed that depression, anxiety, and stress were highly interconnected '
    'and served as central nodes in the psychological network. This finding aligns with the tripartite model '
    'of affect (Clark & Watson, 1991) and suggests that interventions targeting one aspect of distress '
    'may have cascading effects on related symptoms.'
)

discussion_p3 = doc.add_paragraph()
discussion_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
discussion_p3.paragraph_format.first_line_indent = Inches(0.5)
discussion_p3.add_run(
    'Second, Bayesian bootstrap regression identified stress as the strongest predictor of depression, '
    'followed by anxiety. The positive association between self-compassion and depression was unexpected '
    'and may reflect the role of self-focused attention in depressive rumination. Future research should '
    'examine the specific facets of self-compassion that may be adaptive versus maladaptive.'
)

discussion_p4 = doc.add_paragraph()
discussion_p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
discussion_p4.paragraph_format.first_line_indent = Inches(0.5)
discussion_p4.add_run(
    'Third, gender comparison analyses revealed that females reported higher levels of psychological distress, '
    'while males reported higher cyber victimization. These findings are consistent with previous research '
    'documenting gender differences in mental health outcomes and online experiences. The higher rates of '
    'cyber victimization among males may reflect different patterns of online behavior and exposure to risk.'
)

# Limitations
limitations_heading = doc.add_paragraph()
limitations_run = limitations_heading.add_run('Limitations and Future Directions')
limitations_run.bold = True
limitations_run.italic = True

limitations_p = doc.add_paragraph()
limitations_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
limitations_p.paragraph_format.first_line_indent = Inches(0.5)
limitations_p.add_run(
    'Several limitations should be considered when interpreting these findings. First, the cross-sectional '
    'design precludes causal inferences. Second, the reliance on self-report measures may be subject to '
    'response biases. Third, the convenience sampling method limits generalizability. Future research should '
    'employ longitudinal designs and incorporate behavioral measures of social media use.'
)

# Conclusions
conclusions_heading = doc.add_paragraph()
conclusions_run = conclusions_heading.add_run('Conclusions')
conclusions_run.bold = True
conclusions_run.italic = True

conclusions_p = doc.add_paragraph()
conclusions_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
conclusions_p.paragraph_format.first_line_indent = Inches(0.5)
conclusions_p.add_run(
    'The present study contributes to the growing literature on the psychological correlates of social media use. '
    'Network analysis revealed the central role of psychological distress symptoms in the overall network structure, '
    'while Bayesian bootstrap methods provided robust estimates of predictive relationships. These findings have '
    'implications for clinical practice and highlight potential targets for intervention in digitally connected populations.'
)

doc.add_page_break()

# =============================================================================
# REFERENCES
# =============================================================================

references_heading = doc.add_paragraph()
references_run = references_heading.add_run('References')
references_run.bold = True

references = [
    'Clark, L. A., & Watson, D. (1991). Tripartite model of anxiety and depression: Psychometric evidence and taxonomic implications. Journal of Abnormal Psychology, 100(3), 316-336. https://doi.org/10.1037/0021-843X.100.3.316',
    'Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Psychology Foundation of Australia.',
    'Rubin, D. B. (1981). The Bayesian bootstrap. The Annals of Statistics, 9(1), 130-134. https://doi.org/10.1214/aos/1176345338',
]

for ref in references:
    ref_p = doc.add_paragraph()
    ref_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref_p.paragraph_format.left_indent = Inches(0.5)
    ref_p.paragraph_format.first_line_indent = Inches(-0.5)
    ref_p.add_run(ref)

# =============================================================================
# SAVE DOCUMENT
# =============================================================================

output_path = 'data/Analysis_Report_APA7_English.docx'
doc.save(output_path)
print(f"APA 7 format report saved: {output_path}")
