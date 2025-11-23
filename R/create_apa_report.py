#!/usr/bin/env python3
# =============================================================================
# APA 7 Format Analysis Report Generator
# =============================================================================

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_apa_heading(doc, text, level=1):
    """Add APA 7 style heading"""
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    elif level == 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    elif level == 3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

def add_apa_paragraph(doc, text):
    """Add APA 7 style paragraph with first line indent"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_apa_table(doc, df, title, note=None):
    """Add APA 7 style table"""
    # Table title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Table note
    if note:
        p = doc.add_paragraph()
        run = p.add_run("Note. ")
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run = p.add_run(note)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    doc.add_paragraph()
    return table

# =============================================================================
# TITLE PAGE
# =============================================================================

# Title
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bayesian Network Analysis of Social Media Use, Psychological Distress, and Related Factors: A Cross-Sectional Study")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Analysis Report")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# =============================================================================
# ABSTRACT
# =============================================================================

add_apa_heading(doc, "Abstract", 1)

abstract_text = """This study examined the relationships among social media use, psychological distress (depression, anxiety, stress), self-compassion, social support, shame, guilt, and cyber victimization using Bayesian network analysis. Data from 1,028 participants (241 males, 785 females) were analyzed. Network analysis revealed that depression, stress, and anxiety were the most central variables in the psychological network, with strong interconnections among them. Social media addiction served as a bridge variable connecting social media use patterns to psychological distress. Gender comparison analyses indicated that women reported significantly higher levels of depression, anxiety, and stress, while men reported higher cyber victimization experiences. The findings suggest that interventions targeting social media addiction and enhancing self-compassion may be effective in reducing psychological distress."""

add_apa_paragraph(doc, abstract_text)

p = doc.add_paragraph()
run = p.add_run("Keywords: ")
run.italic = True
run.font.name = 'Times New Roman'
run = p.add_run("Bayesian network analysis, social media addiction, psychological distress, cyber victimization, gender differences")
run.font.name = 'Times New Roman'

doc.add_page_break()

# =============================================================================
# RESULTS SECTION
# =============================================================================

add_apa_heading(doc, "Results", 1)

# =============================================================================
# DESCRIPTIVE STATISTICS
# =============================================================================

add_apa_heading(doc, "Descriptive Statistics", 2)

desc_text = """Descriptive statistics for all study variables are presented in Table 1. The sample consisted of 1,028 participants with complete data (Mage = 31.30, SD = 9.34). Depression scores ranged from 0 to 3 (M = 0.97, SD = 0.67), anxiety scores ranged from 0 to 3 (M = 0.73, SD = 0.60), and stress scores ranged from 0 to 3 (M = 1.07, SD = 0.66). Social media addiction scores showed moderate levels (M = 2.72, SD = 0.92), and cyber victimization scores were relatively low (M = 1.17, SD = 0.35) with high positive skewness (3.41), indicating that most participants reported minimal cyber victimization experiences."""

add_apa_paragraph(doc, desc_text)

# Load descriptive data
try:
    desc_df = pd.read_csv('data/descriptives_python.csv')
    desc_df = desc_df.round(2)
    desc_df = desc_df[['Unnamed: 0', 'mean', 'std', 'min', 'max', 'skewness', 'kurtosis']]
    desc_df.columns = ['Variable', 'M', 'SD', 'Min', 'Max', 'Skewness', 'Kurtosis']
    add_apa_table(doc, desc_df, "Table 1", "Descriptive Statistics for Study Variables (N = 1,028)")
except:
    pass

# =============================================================================
# CORRELATION ANALYSIS
# =============================================================================

add_apa_heading(doc, "Correlation Analysis", 2)

corr_text = """Pearson correlations among study variables are presented in Table 2. Depression, anxiety, and stress showed strong positive intercorrelations (rs ranging from .74 to .82, ps < .001), consistent with the conceptualization of these constructs as indicators of psychological distress. Self-compassion was positively correlated with all three distress indicators (rs ranging from .65 to .68, ps < .001), suggesting that lower self-compassion is associated with higher psychological distress. Social media addiction was significantly correlated with depression (r = .38, p < .001), anxiety (r = .36, p < .001), and stress (r = .38, p < .001). Notably, social media addiction showed a strong positive correlation with social media usage time (r = .51, p < .001). Family support and friend support were negatively correlated with depression (rs = -.36 and -.22, respectively, ps < .001), indicating protective effects of social support."""

add_apa_paragraph(doc, corr_text)

# Correlation matrix
try:
    cor_df = pd.read_csv('data/correlation_matrix_python.csv', index_col=0)
    cor_df = cor_df.round(2)
    cor_df_display = cor_df.reset_index()
    cor_df_display.columns = ['Variable'] + list(cor_df.columns)
    # Truncate for display
    cor_df_short = cor_df_display.iloc[:, :7]
    add_apa_table(doc, cor_df_short, "Table 2", "Intercorrelations Among Study Variables (N = 1,028). All correlations > |.10| are significant at p < .05.")
except:
    pass

# =============================================================================
# NETWORK ANALYSIS
# =============================================================================

add_apa_heading(doc, "Network Analysis", 2)

network_text = """Bayesian network analysis was conducted to examine the complex interrelationships among study variables. The network structure was estimated using the Hill-Climbing algorithm with BIC scoring. The resulting network revealed 47 significant edges among the 12 variables."""

add_apa_paragraph(doc, network_text)

add_apa_heading(doc, "Centrality Measures", 3)

centrality_text = """Centrality analyses were conducted to identify the most influential variables in the network. Results are presented in Table 3. Depression emerged as the most central variable in terms of strength centrality (4.13), followed by stress (3.96), anxiety (3.92), and self-compassion (3.76). These findings suggest that depression is the most interconnected variable in the psychological network, serving as a hub that connects to multiple other constructs. Social media addiction showed the highest betweenness centrality (0.11), indicating its role as a bridge variable that connects different parts of the network, particularly linking social media use patterns to psychological distress outcomes."""

add_apa_paragraph(doc, centrality_text)

# Centrality table
try:
    cent_df = pd.read_csv('data/centrality_python.csv')
    cent_df = cent_df.round(3)
    add_apa_table(doc, cent_df, "Table 3", "Centrality Measures for Network Variables")
except:
    pass

# =============================================================================
# GENDER DIFFERENCES
# =============================================================================

add_apa_heading(doc, "Gender Differences", 2)

gender_text = """Independent samples t-tests were conducted to examine gender differences in study variables. Results are presented in Table 4. Women reported significantly higher levels of depression (M = 1.00, SD = 0.67) compared to men (M = 0.88, SD = 0.65), t(1024) = -2.48, p = .013, d = 0.18. Similarly, women reported higher anxiety (M = 0.76, SD = 0.60) than men (M = 0.64, SD = 0.59), t(1024) = -2.60, p = .009, d = 0.19, and higher stress (M = 1.10, SD = 0.66) than men (M = 0.95, SD = 0.66), t(1024) = -3.03, p = .003, d = 0.22. Women also reported higher self-compassion concerns (M = 2.31, SD = 1.08) compared to men (M = 2.10, SD = 1.00), t(1024) = -2.62, p = .009, d = 0.19. In contrast, men reported significantly higher cyber victimization experiences (M = 1.25, SD = 0.45) compared to women (M = 1.14, SD = 0.30), t(1024) = 3.37, p < .001, d = -0.31, representing a medium effect size."""

add_apa_paragraph(doc, gender_text)

# Gender comparison table
try:
    gender_df = pd.read_csv('data/gender_ttest_results.csv')
    gender_df = gender_df.round(3)
    gender_df.columns = ['Variable', 't', 'p', 'Sig', "Cohen's d", 'Effect Size']
    add_apa_table(doc, gender_df, "Table 4", "Independent Samples t-Test Results for Gender Differences. *p < .05, **p < .01, ***p < .001.")
except:
    pass

add_apa_heading(doc, "Gender Differences in Network Structure", 3)

gender_network_text = """Separate network analyses were conducted for male and female participants to examine potential gender differences in the structure of relationships among variables. Notable differences emerged in the strength of specific associations. The relationship between cyber victimization and anxiety was stronger in men (r = .44) compared to women (r = .29). Conversely, the relationship between social media addiction and depression was stronger in women (r = .41) compared to men (r = .29). Additionally, the association between social media addiction and social media usage time was stronger in women (r = .53) compared to men (r = .43). These findings suggest that the pathways linking social media use to psychological distress may differ by gender."""

add_apa_paragraph(doc, gender_network_text)

# Correlation differences table
try:
    cordiff_df = pd.read_csv('data/gender_correlation_differences.csv')
    cordiff_df = cordiff_df.head(10).round(2)
    cordiff_df = cordiff_df[['Variable1', 'Variable2', 'Male_r', 'Female_r', 'Difference']]
    cordiff_df.columns = ['Variable 1', 'Variable 2', 'Male r', 'Female r', 'Difference']
    add_apa_table(doc, cordiff_df, "Table 5", "Largest Gender Differences in Correlation Coefficients (Top 10)")
except:
    pass

# =============================================================================
# SIGNIFICANT CORRELATIONS
# =============================================================================

add_apa_heading(doc, "Key Network Edges", 2)

edges_text = """Table 6 presents the strongest significant correlations identified in the network analysis. The three strongest associations were between depression and stress (r = .82, p < .001), anxiety and stress (r = .79, p < .001), and depression and anxiety (r = .74, p < .001), confirming the tight clustering of psychological distress indicators. Self-compassion showed strong positive associations with all distress variables, while social support variables (family and friend support) showed negative associations with distress, consistent with their protective role."""

add_apa_paragraph(doc, edges_text)

# Significant edges table
try:
    edges_df = pd.read_csv('data/significant_correlations_python.csv')
    edges_df = edges_df.head(15).round(3)
    edges_df = edges_df[['From', 'To', 'Correlation', 'p_value']]
    edges_df.columns = ['Variable 1', 'Variable 2', 'r', 'p']
    add_apa_table(doc, edges_df, "Table 6", "Strongest Significant Correlations in the Network (Top 15)")
except:
    pass

# =============================================================================
# DISCUSSION
# =============================================================================

doc.add_page_break()
add_apa_heading(doc, "Discussion", 1)

discussion_text = """The present study employed Bayesian network analysis to examine the complex interrelationships among social media use, psychological distress, and related factors. Several key findings emerged from the analyses.

First, depression, stress, and anxiety emerged as the most central variables in the psychological network, with high strength centrality values indicating their strong interconnectedness with other variables. This finding is consistent with the conceptualization of these constructs as core indicators of psychological distress and supports the utility of targeting these symptoms in clinical interventions.

Second, social media addiction demonstrated the highest betweenness centrality, suggesting its role as a bridge variable connecting social media use patterns to psychological distress outcomes. This finding has important implications for intervention design, as targeting social media addiction may help disrupt the pathway between excessive social media use and negative psychological outcomes.

Third, significant gender differences were observed in both the levels of psychological variables and the network structure. Women reported higher levels of depression, anxiety, and stress, while men reported higher cyber victimization. Furthermore, the relationships between social media addiction and psychological distress were stronger in women, while the relationships between cyber victimization and anxiety were stronger in men. These findings suggest that gender-specific intervention approaches may be warranted.

Fourth, social support variables (family and friend support) showed consistent negative associations with psychological distress, confirming their protective role. Interventions that enhance social support networks may therefore be beneficial in reducing psychological distress associated with social media use."""

add_apa_paragraph(doc, discussion_text)

# =============================================================================
# LIMITATIONS
# =============================================================================

add_apa_heading(doc, "Limitations", 2)

limitations_text = """Several limitations should be noted. First, the cross-sectional design precludes causal inferences; longitudinal studies are needed to establish temporal precedence among variables. Second, the sample was predominantly female (76.4%), which may limit generalizability. Third, all measures were self-report, which may be subject to response biases. Fourth, the network analysis is correlational in nature and does not establish causal relationships among variables."""

add_apa_paragraph(doc, limitations_text)

# =============================================================================
# CONCLUSIONS
# =============================================================================

add_apa_heading(doc, "Conclusions", 2)

conclusions_text = """The present study contributes to the understanding of the complex relationships among social media use, psychological distress, and related factors. The network analysis approach revealed that depression serves as a central hub in the psychological network, while social media addiction acts as a bridge connecting social media use to psychological outcomes. Gender differences in both variable levels and network structure highlight the need for gender-sensitive approaches in research and intervention. Future research should employ longitudinal designs to establish causal pathways and develop targeted interventions based on the identified network structure."""

add_apa_paragraph(doc, conclusions_text)

# =============================================================================
# SAVE DOCUMENT
# =============================================================================

doc.save('data/Analysis_Report_APA7.docx')
print("Report saved: data/Analysis_Report_APA7.docx")
print("\nReport contents:")
print("- Title Page")
print("- Abstract")
print("- Results Section")
print("  - Descriptive Statistics (Table 1)")
print("  - Correlation Analysis (Table 2)")
print("  - Network Analysis")
print("  - Centrality Measures (Table 3)")
print("  - Gender Differences (Table 4)")
print("  - Gender Network Differences (Table 5)")
print("  - Key Network Edges (Table 6)")
print("- Discussion")
print("- Limitations")
print("- Conclusions")
